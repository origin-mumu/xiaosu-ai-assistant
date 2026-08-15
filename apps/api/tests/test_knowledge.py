from io import BytesIO

import pymupdf
from docx import Document

from xiaosu.knowledge.chunker import chunk_segments
from xiaosu.knowledge.parser import parse_document


def test_markdown_parser_preserves_section_and_paragraph() -> None:
    segments = parse_document(
        "handbook.md",
        "# 请假制度\n\n员工应提前申请。\n\n## 审批\n\n由直属主管审批。".encode(),
    )

    assert [segment.section_title for segment in segments] == ["请假制度", "审批"]
    assert [segment.paragraph_number for segment in segments] == [1, 2]


def test_docx_parser_preserves_heading() -> None:
    document = Document()
    document.add_heading("报销制度", level=1)
    document.add_paragraph("发票应真实有效。")
    buffer = BytesIO()
    document.save(buffer)

    segments = parse_document("policy.docx", buffer.getvalue())

    assert segments[0].section_title == "报销制度"
    assert segments[0].text == "发票应真实有效。"


def test_chunker_splits_long_text_with_overlap() -> None:
    segments = parse_document("notes.txt", ("一二三四五六七八九十" * 20).encode())
    chunks = chunk_segments(segments, chunk_size=50, overlap=10)

    assert len(chunks) > 1
    assert chunks[0].content[-10:] == chunks[1].content[:10]


def test_pdf_parser_preserves_page_number() -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Expense policy requires a valid invoice.")
    content = document.tobytes()
    document.close()

    segments = parse_document("policy.pdf", content)

    assert segments[0].page_number == 1
    assert "valid invoice" in segments[0].text

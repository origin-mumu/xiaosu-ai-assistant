from __future__ import annotations

import io
import random
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "data" / "documents"
BLUE = RGBColor(74, 93, 230)
DARK_BLUE = RGBColor(25, 39, 75)
MUTED = RGBColor(105, 118, 145)


EMPLOYEE_HANDBOOK = """# 员工手册

> 适用范围：全体正式员工、试用期员工及经批准的长期派驻人员。制度更新时间：2026 年 8 月。

## 1. 工作时间与考勤

- 标准工作时间为周一至周五 09:00—18:00，午休时间为 12:00—13:00。
- 员工应通过钉钉完成上下班打卡。迟到、早退或漏卡需在 2 个工作日内提交补卡申请，每人每月原则上不超过 3 次。
- 外出拜访、出差或远程办公须事先提交申请并经直属主管批准；审批记录可作为当日考勤依据。
- 连续旷工 2 个工作日，或年度累计旷工 3 个工作日，将按严重违纪流程处理。

## 2. 年假与其他休假

- 累计工作满 1 年不满 10 年，法定年假 5 天；满 10 年不满 20 年，法定年假 10 天；满 20 年及以上，法定年假 15 天。
- 新入职员工当年度年假按剩余日历天数折算，不足 0.5 天不计，满 0.5 天按 1 天计。
- 年假应至少提前 3 个工作日申请；连续休假 3 天及以上，建议提前 7 个工作日申请。
- 当年度未休年假可结转至次年 3 月 31 日，逾期未使用且公司已安排休假的部分按制度处理。
- 事假按小时申请，最小单位 1 小时；病假应在返岗后 3 个工作日内补交有效医疗证明。
- 婚假、产假、陪产假、育儿假及丧假按员工工作地现行法规和公司流程执行。

## 3. 加班与调休

- 加班必须事前审批，未经审批的延时工作不计入有效加班。
- 工作日加班连续满 1 小时、休息日加班连续满 2 小时方可登记；用餐和通勤时间不计入加班。
- 休息日加班优先安排调休，调休原则上在 3 个月内使用；法定节假日加班依法支付加班工资。

## 4. 出差管理

- 员工应至少提前 3 个工作日发起出差申请，填写目的地、事由、日期、预算和同行人。
- 国内交通优先选择高铁二等座或经济舱；一线城市住宿上限 600 元/晚，其他城市 450 元/晚。
- 超标准预订须在出行前取得部门负责人书面批准；紧急情况应在返程后 2 个工作日内补充说明。

## 5. 报销制度

- 日常费用应在发生后 30 个自然日内提交；差旅费用应在返程后 10 个工作日内提交。
- 报销材料包括审批单、合法发票、消费明细和支付凭证。电子发票须上传原始 PDF，不得提交截图。
- 单笔 5,000 元及以上费用还需成本中心负责人复核；超标准费用须附书面说明。
- 财务每周二、周四集中付款，审批完成且材料齐全的报销通常在 5 个工作日内到账。

## 6. 薪酬与福利

- 工资于每月 10 日发放，遇法定节假日原则上提前。工资单可在薪酬系统查看。
- 公司依法缴纳社会保险和住房公积金，并提供年度体检、节日福利及符合条件的通讯补贴。

## 7. 信息安全与保密

- 客户数据、源代码、访问密钥、未公开经营数据不得通过个人邮箱、公共网盘或未经批准的 AI 服务传输。
- 离开工位须锁屏；发现账号异常、设备遗失或疑似泄密，应立即联系 IT 与信息安全负责人。

## 8. 咨询与申诉

制度问题可先咨询直属主管或 HR 服务台；对考勤、薪酬、报销结果有异议的，应在收到结果后 5 个工作日内提交申诉。
"""


ONBOARDING_GUIDE = """# 新人入职指南

> 本指南帮助新员工在入职前 10 个工作日内完成账号、设备、培训和团队融入。

## 入职前准备

- 按 Offer 邮件要求提交身份证明、学历证明、银行卡和证件照；敏感材料仅上传至 HR 指定系统。
- 在入职日前完成电子合同预览、紧急联系人填写和个税专项附加扣除确认。

## 入职第一天

1. 09:30 前往前台签到，完成劳动合同签署和员工证领取。
2. 由行政发放办公设备，员工核对资产编号、配置和外观后签收。
3. 直属主管介绍团队、岗位职责、试用期目标和协作方式。
4. 完成信息安全、消防安全和员工制度必修课程。

## 账号开通

- 企业邮箱、钉钉、OA、知识库和代码平台基础账号由 IT 自动创建，初始信息发送到入职登记手机。
- 首次登录必须修改密码并绑定 MFA。生产环境、财务系统和客户数据权限需单独提交工单审批。
- 账号在入职日 12:00 前仍不可用时，联系 IT 服务台，并提供姓名、员工编号和系统名称。

## 设备申领

- 标准配置包括笔记本电脑、电源适配器、鼠标和电脑包；研发岗位可申请一台外接显示器。
- 特殊配置须由直属主管说明业务必要性，经部门负责人和 IT 资产管理员批准。
- 设备故障先提交 IT 工单，不得自行拆机或委托外部维修。

## 培训流程

- 入职 3 个工作日内完成公司文化、制度与信息安全课程。
- 入职 5 个工作日内完成岗位基础课程和工具培训。
- 入职 10 个工作日内由导师组织第一次学习回顾，并记录待补课程。

## 30/60/90 天融入计划

- 30 天：熟悉业务、流程和关键联系人，完成第一个独立任务。
- 60 天：能够独立承担常规工作，并提出至少一项流程改进建议。
- 90 天：完成试用期总结，由直属主管发起转正评估。

## 常见联系人

- 合同、社保、请假：HR 服务台。
- 账号、网络、设备：IT 服务台。
- 门禁、工位、会议室：行政服务台。
- 报销、发票、付款：财务共享中心。
"""


FAQ = """# 员工常见问题 FAQ

## 休假与考勤

1. **我有几天年假？** 累计工作 1—10 年为 5 天，10—20 年为 10 天，20 年及以上为 15 天；新员工按当年剩余日历天数折算。
2. **年假可以结转吗？** 可以，最晚应在次年 3 月 31 日前使用。
3. **忘记打卡怎么办？** 在 2 个工作日内从钉钉提交补卡申请，每月原则上不超过 3 次。
4. **远程办公算出勤吗？** 事前申请并经直属主管批准后，审批记录可作为考勤依据。
5. **病假需要证明吗？** 需要，返岗后 3 个工作日内补交有效医疗证明。
6. **事假最小申请单位是什么？** 1 小时。
7. **加班必须审批吗？** 必须事前审批，未经审批不计入有效加班。
8. **调休多久内用完？** 原则上在加班发生后 3 个月内使用。

## 报销与出差

9. **日常费用多久内报销？** 费用发生后 30 个自然日内。
10. **差旅费用多久内报销？** 返程后 10 个工作日内。
11. **报销要什么材料？** 审批单、合法发票、消费明细和支付凭证。
12. **电子发票能传截图吗？** 不能，应上传原始 PDF。
13. **什么时候到账？** 材料齐全并审批完成后通常 5 个工作日内到账。
14. **出差要提前多久申请？** 至少提前 3 个工作日。
15. **高铁和飞机标准是什么？** 原则上高铁二等座或经济舱。
16. **住宿上限是多少？** 一线城市 600 元/晚，其他城市 450 元/晚。
17. **费用超标怎么办？** 出行或消费前取得负责人批准，并在报销时附说明。

## 入职、账号与设备

18. **入职第一天几点到？** 09:30 前到前台签到。
19. **邮箱和钉钉谁开通？** IT 自动创建，初始信息发送至入职登记手机。
20. **为什么生产系统不能登录？** 生产权限需额外工单审批，不随基础账号自动开通。
21. **首次登录要做什么？** 修改初始密码并绑定 MFA。
22. **电脑坏了怎么办？** 提交 IT 工单，不要自行拆机或外修。
23. **可以申请显示器吗？** 研发岗位可按标准申请一台，其他岗位说明业务需要后审批。
24. **设备丢失怎么办？** 立即联系直属主管、IT 和信息安全负责人。
25. **必修培训多久完成？** 公司制度与安全课程 3 个工作日内，岗位课程 5 个工作日内。

## 办公与行政

26. **如何预订会议室？** 在钉钉会议室应用选择时间，填写主题和参会人数。
27. **会议室没人使用能占用吗？** 不建议，应先完成预订；临时使用不得影响已预约会议。
28. **访客如何入园？** 接待人至少提前 2 小时提交访客预约，访客凭有效证件登记。
29. **工牌忘带怎么办？** 到前台领取临时访客卡，当日离开前归还。
30. **办公用品如何领取？** 常规用品在行政自助区登记领取，大额或特殊物品需审批。

## 薪酬、成长与离职

31. **工资哪天发？** 每月 10 日，遇法定节假日原则上提前。
32. **工资单在哪里看？** 登录薪酬系统，在“我的工资单”中查看。
33. **如何申请培训？** 在学习平台选课；付费外训需主管与部门负责人审批。
34. **试用期如何转正？** 入职约 90 天完成总结，由直属主管发起转正评估。
35. **离职要归还什么？** 员工证、电脑、配件、门禁卡及其他登记资产。
36. **离职账号何时关闭？** 原则上在最后工作日结束后关闭，交接所需例外须审批。

## 安全与数据

37. **能把文件发到个人邮箱吗？** 不能，内部或敏感资料必须通过公司批准的渠道传输。
38. **能把客户数据发给公共 AI 吗？** 不能，除非该服务已通过公司安全与合规审批。
39. **收到可疑邮件怎么办？** 不点击链接或附件，使用“举报钓鱼”功能并通知信息安全团队。
40. **访问密钥可以写在代码里吗？** 不能，应存入公司密钥管理系统并按最小权限使用。
"""


SECURITY_GUIDE = """信息安全规范
版本：2026.08｜适用对象：全体员工及外包人员

一、账号与密码
公司账号仅限本人使用。密码不得复用，首次登录必须修改初始密码并启用 MFA。发现异常登录应立即重置密码并报备。

二、权限管理
权限遵循最小必要原则。生产、财务、客户数据等敏感系统必须走工单审批，每季度由系统负责人复核。

三、终端安全
办公电脑必须启用磁盘加密、自动锁屏和终端防护，不得擅自关闭安全软件、越狱、刷机或安装来源不明的软件。

四、数据分级
数据分为公开、内部、机密、严格机密四级。客户数据、源代码、密钥、未公开经营数据至少按机密级管理。

五、文件传输
内部及敏感文件仅能通过企业网盘、受控邮件或批准的协作系统传输，不得发送至个人邮箱或公共网盘。

六、研发安全
访问密钥不得硬编码或提交 Git 仓库，应存入密钥管理服务。代码合并前完成审查，严重漏洞必须在发布前修复。

七、生成式 AI 使用
未获批准的 AI 服务不得接收客户数据、员工隐私、源代码、合同或未公开经营信息；输出内容必须由员工复核。

八、钓鱼与社工
不扫描陌生二维码，不向来电者提供验证码。可疑邮件使用“举报钓鱼”功能并通知信息安全团队。

九、安全事件
设备遗失、疑似泄密、恶意软件或账号被盗应立即联系 IT 和信息安全负责人，保留现场和日志，不得自行隐瞒处理。

十、离职与交接
最后工作日归还全部资产，移交业务文件和账号所有权；个人不得复制公司数据，系统权限原则上当日关闭。
"""


VISITOR_GUIDE = """会议室与访客指南
更新时间：2026 年 8 月

1. 会议室预订
在钉钉会议室应用中选择地点和时间，填写会议主题、人数与设备需求。超过 30 分钟未签到，系统可自动释放会议室。

2. 使用规范
按预约时间使用，会议结束后关闭显示器、空调和照明，带走个人物品并恢复桌椅。视频设备故障请联系行政服务台。

3. 取消与变更
计划变化时应及时取消预订。连续三次预订未使用且未取消，行政可暂停预订权限 7 天。

4. 访客预约
接待人至少提前 2 小时在钉钉提交访客预约，填写姓名、手机号、来访时间、事由和接待区域。

5. 访客入园
访客凭有效身份证件在前台登记并领取访客证。接待人应全程陪同，访客不得进入未授权区域。

6. 保密要求
涉及客户、产品路线或财务信息的会议应关闭门窗并及时清理白板。访客拍照、录音或接入内网需事先批准。

7. 临时工牌
员工忘带工牌可在前台登记领取临时卡，卡片仅限当日使用，离开办公区前必须归还。

8. 紧急情况
发现人员受伤、火警或其他紧急情况时，拨打园区应急电话并通知前台，按照疏散指引前往集合点。
"""


TRAVEL_SECTIONS = {
    "适用范围与原则": [
        "本制度适用于因公发生的国内出差、交通、住宿、餐饮和必要业务招待费用。",
        "员工应坚持真实、必要、节约原则，不得拆分费用、重复报销或使用与业务无关的票据。",
    ],
    "出差申请": [
        "至少提前 3 个工作日提交申请，写明目的地、事由、日期、预算和同行人。",
        "未经审批产生的费用原则上不予报销；紧急出差应在返程后 2 个工作日内补充说明。",
    ],
    "交通与住宿标准": [
        "国内交通优先高铁二等座或经济舱；市内交通选择公共交通或合规网约车。",
        "一线城市住宿上限为 600 元/晚，其他城市 450 元/晚；会议统一住宿按会议标准执行。",
        "确需升级标准时，必须在预订前取得部门负责人书面批准。",
    ],
    "补贴与业务招待": [
        "差旅餐饮补贴为每人每天 100 元，客户或会议承担餐饮的日期不重复领取。",
        "业务招待应事前审批，报销时填写参与人、单位、业务事由和人均金额。",
    ],
    "报销材料": [
        "提交差旅审批单、合法发票、消费明细、支付凭证和必要的行程证明。",
        "电子发票必须上传原始 PDF；抬头和税号错误的发票应先更正。",
    ],
    "时间与付款": [
        "返程后 10 个工作日内提交报销，逾期需附原因并由部门负责人确认。",
        "审批完成且材料齐全后，财务通常在 5 个工作日内付款。",
    ],
    "禁止事项与咨询": [
        "严禁虚假行程、虚开发票、重复报销、私费公报或替他人报销。",
        "对标准或单据有疑问时，应在消费前咨询财务共享中心。",
    ],
}


IT_PDF_SECTIONS = {
    "标准设备与签收": "标准办公包包含笔记本电脑、电源适配器、鼠标和电脑包。研发岗位可申请一台外接显示器。员工签收时应核对资产编号、型号、配件和外观，并在资产系统确认。",
    "账号和多因素认证": "企业邮箱、钉钉、OA 和知识库基础账号由 IT 创建。首次登录必须修改初始密码并绑定 MFA。生产、财务和客户数据权限需要单独提交工单审批。",
    "软件安装与管理员权限": "软件应从公司软件中心安装。需要本地管理员权限时，应说明用途和时限，由直属主管与 IT 审批；不得安装破解软件或关闭安全防护。",
    "故障报修": "设备故障应提交 IT 工单，说明资产编号、故障现象和紧急程度。不得自行拆机或委托外部维修；如怀疑数据丢失或安全事件，应先断网并保留现场。",
    "外借与异地使用": "设备外借须在资产系统登记借用人和归还日期。携带设备出差时不得托运，公共场所应避免屏幕泄露，离开时随身携带或锁入安全柜。",
    "归还与数据清理": "岗位变更或离职时，应在最后工作日归还全部设备和配件。IT 完成数据备份确认后统一清理设备，员工不得自行复制公司数据。",
}


SALES_PDF_SECTIONS = {
    "客户信息登记": "新增客户应在 CRM 记录公司名称、联系人、来源、负责人和跟进状态。不得在个人通讯录或未批准的表格中长期保存客户敏感信息。",
    "商机阶段定义": "商机依次进入线索确认、需求分析、方案报价、商务谈判、赢单或丢单阶段。负责人应在关键活动后 1 个工作日内更新阶段、金额和预计成交日。",
    "报价与折扣": "报价须使用审批后的标准模板。低于标准价格 10% 的折扣需销售负责人审批，20% 及以上还需财务和业务负责人批准。",
    "合同与回款": "合同必须经过法务和财务审核后方可签署。订单状态以订单系统为准；销售人员应跟踪开票、回款和逾期情况，不得私自承诺账期。",
    "客户投诉处理": "收到投诉后 2 小时内响应，4 小时内确定责任人，重大投诉应立即升级。所有沟通、临时方案和最终结论应记录在工单系统。",
    "数据统计口径": "销售额按订单系统中已生效且未取消订单的含税金额统计。退款在退款生效日冲减；查询周销售额时以公司时区周一 00:00 至周日 23:59 为周期。",
}


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_east_asia_font(run, name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _banner_jpeg() -> bytes:
    width, height = 2200, 760
    randomizer = random.Random(20260815)
    samples = bytearray(width * height * 3)
    for y in range(height):
        ratio = y / max(height - 1, 1)
        for x in range(width):
            wave = 18 if (x // 170 + y // 95) % 2 else 0
            noise = randomizer.randrange(-11, 12)
            index = (y * width + x) * 3
            samples[index] = max(0, min(255, int(90 + 70 * ratio + wave + noise)))
            samples[index + 1] = max(0, min(255, int(108 + 82 * ratio + noise)))
            samples[index + 2] = max(0, min(255, int(228 + 18 * ratio + noise // 2)))
    image = Image.frombytes("RGB", (width, height), bytes(samples))
    output = io.BytesIO()
    image.save(output, format="JPEG", quality=91, optimize=True)
    return output.getvalue()


def write_text_documents() -> None:
    documents = {
        "员工手册.md": EMPLOYEE_HANDBOOK,
        "新人入职指南.md": ONBOARDING_GUIDE,
        "常见问题FAQ.md": FAQ,
        "信息安全规范.txt": SECURITY_GUIDE,
        "会议室与访客指南.txt": VISITOR_GUIDE,
    }
    for filename, content in documents.items():
        (OUTPUT / filename).write_text(content.strip() + "\n", encoding="utf-8")


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("小苏企业知识库 · 内部资料 · 2026")
    _set_east_asia_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def generate_docx(banner: bytes) -> None:
    document = Document()
    configure_docx(document)
    document.add_picture(io.BytesIO(banner), width=Inches(6.7))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("差旅与报销制度")
    _set_east_asia_font(title_run)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("审批、标准、单据和付款的一站式参考")
    _set_east_asia_font(subtitle_run)
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    labels = (
        "适用范围\n全体员工",
        "提交时限\n返程后 10 个工作日",
        "咨询渠道\n财务共享中心",
    )
    for cell, label in zip(table.rows[0].cells, labels):
        _set_cell_shading(cell, "EEF1FF")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        _set_east_asia_font(run)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE

    for title_text, bullets in TRAVEL_SECTIONS.items():
        document.add_heading(title_text, level=1)
        for bullet in bullets:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(bullet)
            _set_east_asia_font(run)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    run = note.add_run(
        "提示：本资料用于知识库检索测试；若与最新公告冲突，以公司正式发布版本为准。"
    )
    _set_east_asia_font(run)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED
    document.save(OUTPUT / "差旅与报销制度.docx")


def register_pdf_font() -> str:
    candidates = (
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    )
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("XiaosuCJK", str(candidate), subfontIndex=0))
            return "XiaosuCJK"
    raise RuntimeError("未找到可用于生成中文 PDF 的字体")


def draw_wrapped_text(
    pdf: canvas.Canvas,
    text: str,
    *,
    x: float,
    y: float,
    width: float,
    font: str,
    size: float,
    leading: float,
) -> float:
    line = ""
    lines: list[str] = []
    for character in text:
        candidate = line + character
        if pdfmetrics.stringWidth(candidate, font, size) <= width:
            line = candidate
        else:
            lines.append(line)
            line = character
    if line:
        lines.append(line)
    pdf.setFont(font, size)
    for content in lines:
        pdf.drawString(x, y, content)
        y -= leading
    return y


def generate_pdf(
    filename: str, title: str, subtitle: str, sections: dict[str, str], banner: bytes
) -> None:
    font = register_pdf_font()
    pdf = canvas.Canvas(str(OUTPUT / filename), pagesize=letter, pageCompression=1)
    pdf.setTitle(title)
    pdf.setAuthor("小苏企业知识库")
    pdf.setSubject(subtitle)
    image = ImageReader(io.BytesIO(banner))

    pdf.drawImage(
        image, 42, 568, width=528, height=182, preserveAspectRatio=False, mask="auto"
    )
    pdf.setFillColor(HexColor("#F7F8FF"))
    pdf.setStrokeColor(HexColor("#E5E8F5"))
    pdf.roundRect(42, 142, 528, 382, 16, stroke=1, fill=1)
    pdf.setFillColor(HexColor("#19274B"))
    pdf.setFont(font, 25)
    pdf.drawString(72, 450, title)
    pdf.setFillColor(HexColor("#59698F"))
    pdf.setFont(font, 13)
    pdf.drawString(72, 405, subtitle)
    pdf.setFont(font, 10)
    pdf.drawString(72, 208, "小苏企业知识库｜内部参考资料｜2026 年 8 月")
    pdf.showPage()

    for page_number, (heading, body) in enumerate(sections.items(), start=2):
        pdf.setFillColor(HexColor("#4A5DE6"))
        pdf.roundRect(42, 680, 528, 70, 14, stroke=0, fill=1)
        pdf.setFillColor(HexColor("#FFFFFF"))
        pdf.setFont(font, 18)
        pdf.drawString(66, 706, heading)
        pdf.setFillColor(HexColor("#1F2B49"))
        draw_wrapped_text(
            pdf, body, x=66, y=620, width=480, font=font, size=12.5, leading=25
        )
        pdf.setStrokeColor(HexColor("#DCE0EC"))
        pdf.line(66, 96, 546, 96)
        pdf.setFillColor(HexColor("#7B879F"))
        pdf.setFont(font, 8.5)
        pdf.drawString(66, 74, "小苏企业知识库 · 可检索原文")
        pdf.drawRightString(546, 74, str(page_number))
        pdf.showPage()

    pdf.save()


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    banner = _banner_jpeg()
    write_text_documents()
    generate_docx(banner)
    generate_pdf(
        "IT设备与账号管理指南.pdf",
        "IT 设备与账号管理指南",
        "设备签收、账号权限、报修与归还流程",
        IT_PDF_SECTIONS,
        banner,
    )
    generate_pdf(
        "销售与客户服务规范.pdf",
        "销售与客户服务规范",
        "CRM、报价、合同、投诉与销售统计口径",
        SALES_PDF_SECTIONS,
        banner,
    )

    legacy_pdf = OUTPUT / "office-equipment-guide.pdf"
    if legacy_pdf.exists():
        legacy_pdf.unlink()

    files = sorted(path for path in OUTPUT.iterdir() if path.is_file())
    total_size = sum(path.stat().st_size for path in files) / (1024 * 1024)
    print(f"Generated {len(files)} sample documents in {OUTPUT} ({total_size:.2f} MiB)")


if __name__ == "__main__":
    main()
